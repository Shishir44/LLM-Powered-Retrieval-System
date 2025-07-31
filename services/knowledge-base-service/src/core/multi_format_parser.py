"""
Multi-Format Document Parser
Unified document parsing for PDF, DOCX, HTML, CSV, XML with metadata extraction
"""

from typing import Dict, Any, List, Optional, Union
import logging
from pathlib import Path
from dataclasses import dataclass
from datetime import datetime
import hashlib

# Document parsing libraries
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import pandas as pd
import xml.etree.ElementTree as ET
import chardet
import mimetypes
import json

@dataclass
class ParsedDocument:
    """Structured representation of a parsed document."""
    content: str
    title: str
    metadata: Dict[str, Any]
    structure: Dict[str, Any]
    images: List[Dict[str, Any]]
    tables: List[Dict[str, Any]]
    links: List[Dict[str, str]]
    raw_format: str
    encoding: str
    parse_errors: List[str]

class MultiFormatParser:
    """Unified parser for multiple document formats with structure preservation."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
        # Supported formats and their handlers
        self.format_handlers = {
            'pdf': self._parse_pdf,
            'docx': self._parse_docx,
            'html': self._parse_html,
            'htm': self._parse_html,
            'csv': self._parse_csv,
            'xml': self._parse_xml,
            'txt': self._parse_text,
            'json': self._parse_json,
            'md': self._parse_markdown
        }
        
        # MIME type mappings
        self.mime_mappings = {
            'application/pdf': 'pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'text/html': 'html',
            'text/csv': 'csv',
            'application/xml': 'xml',
            'text/xml': 'xml',
            'text/plain': 'txt',
            'application/json': 'json',
            'text/markdown': 'md'
        }

    async def parse_document(self, 
                           file_path: Union[str, Path], 
                           content: Optional[bytes] = None,
                           format_hint: Optional[str] = None) -> ParsedDocument:
        """Parse document from file path or content bytes."""
        
        file_path = Path(file_path) if isinstance(file_path, str) else file_path
        
        try:
            # Detect format
            doc_format = self._detect_format(file_path, content, format_hint)
            
            # Detect encoding for text-based formats
            encoding = self._detect_encoding(file_path, content) if content else 'utf-8'
            
            # Get appropriate handler
            handler = self.format_handlers.get(doc_format.lower())
            if not handler:
                raise ValueError(f"Unsupported format: {doc_format}")
            
            # Parse document
            self.logger.info(f"Parsing {doc_format.upper()} document: {file_path.name}")
            parsed_doc = await handler(file_path, content, encoding)
            
            # Add common metadata
            parsed_doc.metadata.update({
                'file_path': str(file_path),
                'file_name': file_path.name,
                'file_size': file_path.stat().st_size if file_path.exists() else len(content or b''),
                'format': doc_format,
                'parsed_at': datetime.now().isoformat(),
                'parser_version': '1.0.0'
            })
            
            return parsed_doc
            
        except Exception as e:
            self.logger.error(f"Failed to parse document {file_path}: {e}")
            return ParsedDocument(
                content="",
                title=file_path.name,
                metadata={'error': str(e)},
                structure={},
                images=[],
                tables=[],
                links=[],
                raw_format=format_hint or 'unknown',
                encoding='utf-8',
                parse_errors=[str(e)]
            )

    def _detect_format(self, file_path: Path, content: Optional[bytes], hint: Optional[str]) -> str:
        """Detect document format from multiple sources."""
        
        # Use hint if provided
        if hint and hint.lower() in self.format_handlers:
            return hint.lower()
        
        # Check file extension
        ext = file_path.suffix.lower().lstrip('.')
        if ext in self.format_handlers:
            return ext
        
        # Check MIME type
        if content:
            mime_type, _ = mimetypes.guess_type(file_path)
            if mime_type in self.mime_mappings:
                return self.mime_mappings[mime_type]
        
        # Default to text
        return 'txt'

    def _detect_encoding(self, file_path: Path, content: Optional[bytes]) -> str:
        """Detect text encoding."""
        try:
            if content:
                result = chardet.detect(content)
                return result.get('encoding', 'utf-8') or 'utf-8'
            elif file_path.exists():
                with open(file_path, 'rb') as f:
                    raw_data = f.read(10000)  # Read first 10KB
                    result = chardet.detect(raw_data)
                    return result.get('encoding', 'utf-8') or 'utf-8'
        except Exception:
            pass
        return 'utf-8'

    async def _parse_pdf(self, file_path: Path, content: Optional[bytes], encoding: str) -> ParsedDocument:
        """Parse PDF documents with text, tables, and metadata extraction."""
        
        text_content = []
        tables = []
        images = []
        metadata = {}
        structure = {'pages': [], 'sections': []}
        errors = []
        
        try:
            # Use pdfplumber for better text and table extraction
            if content:
                import io
                pdf_file = io.BytesIO(content)
            else:
                pdf_file = str(file_path)
            
            with pdfplumber.open(pdf_file) as pdf:
                # Extract metadata
                if pdf.metadata:
                    metadata.update({
                        'title': pdf.metadata.get('Title', ''),
                        'author': pdf.metadata.get('Author', ''),
                        'subject': pdf.metadata.get('Subject', ''),
                        'creator': pdf.metadata.get('Creator', ''),
                        'creation_date': pdf.metadata.get('CreationDate', ''),
                        'modification_date': pdf.metadata.get('ModDate', ''),
                        'pages': len(pdf.pages)
                    })
                
                # Extract content page by page
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text() or ""
                    text_content.append(page_text)
                    
                    # Extract tables from page
                    page_tables = page.extract_tables()
                    for table_idx, table in enumerate(page_tables):
                        if table:
                            tables.append({
                                'page': page_num,
                                'table_index': table_idx,
                                'data': table,
                                'headers': table[0] if table else [],
                                'rows': table[1:] if len(table) > 1 else []
                            })
                    
                    # Track page structure
                    structure['pages'].append({
                        'page_number': page_num,
                        'text_length': len(page_text),
                        'tables_count': len(page_tables),
                        'bbox': page.bbox
                    })
        
        except Exception as e:
            errors.append(f"PDF parsing error: {e}")
            # Fallback to PyPDF2
            try:
                if content:
                    import io
                    pdf_file = io.BytesIO(content)
                else:
                    pdf_file = open(file_path, 'rb')
                
                reader = PyPDF2.PdfReader(pdf_file)
                for page in reader.pages:
                    text_content.append(page.extract_text())
                
                if not content:
                    pdf_file.close()
                    
            except Exception as e2:
                errors.append(f"Fallback PDF parsing error: {e2}")
        
        # Combine all text
        full_text = '\n\n'.join(text_content)
        title = metadata.get('title') or file_path.stem
        
        return ParsedDocument(
            content=full_text,
            title=title,
            metadata=metadata,
            structure=structure,
            images=images,
            tables=tables,
            links=[],
            raw_format='pdf',
            encoding=encoding,
            parse_errors=errors
        )

    async def _parse_docx(self, file_path: Path, content: Optional[bytes], encoding: str) -> ParsedDocument:
        """Parse DOCX documents with structure preservation."""
        
        text_content = []
        tables = []
        images = []
        structure = {'paragraphs': [], 'sections': [], 'styles': []}
        errors = []
        
        try:
            if content:
                import io
                doc_file = io.BytesIO(content)
            else:
                doc_file = str(file_path)
            
            doc = DocxDocument(doc_file)
            
            # Extract core properties
            metadata = {
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'subject': doc.core_properties.subject or '',
                'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                'modified': str(doc.core_properties.modified) if doc.core_properties.modified else '',
                'paragraphs_count': len(doc.paragraphs),
                'tables_count': len(doc.tables)
            }
            
            # Extract paragraphs with style information
            for para_idx, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
                    structure['paragraphs'].append({
                        'index': para_idx,
                        'text': paragraph.text,
                        'style': paragraph.style.name if paragraph.style else 'Normal',
                        'alignment': str(paragraph.alignment) if paragraph.alignment else None
                    })
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                tables.append({
                    'table_index': table_idx,
                    'data': table_data,
                    'headers': table_data[0] if table_data else [],
                    'rows': table_data[1:] if len(table_data) > 1 else [],
                    'rows_count': len(table.rows),
                    'cols_count': len(table.columns)
                })
            
            # Extract images (basic info)
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    images.append({
                        'type': 'embedded_image',
                        'target': rel.target_ref,
                        'relationship_id': rel.rId
                    })
        
        except Exception as e:
            errors.append(f"DOCX parsing error: {e}")
        
        full_text = '\n\n'.join(text_content)
        title = metadata.get('title') or file_path.stem
        
        return ParsedDocument(
            content=full_text,
            title=title,
            metadata=metadata,
            structure=structure,
            images=images,
            tables=tables,
            links=[],
            raw_format='docx',
            encoding=encoding,
            parse_errors=errors
        )

    async def _parse_html(self, file_path: Path, content: Optional[bytes], encoding: str) -> ParsedDocument:
        """Parse HTML documents with structure and link extraction."""
        
        try:
            if content:
                html_content = content.decode(encoding)
            else:
                with open(file_path, 'r', encoding=encoding) as f:
                    html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Extract metadata
            title = soup.find('title')
            title_text = title.get_text().strip() if title else file_path.stem
            
            meta_tags = soup.find_all('meta')
            metadata = {'title': title_text}
            
            for meta in meta_tags:
                name = meta.get('name') or meta.get('property')
                content_attr = meta.get('content')
                if name and content_attr:
                    metadata[name] = content_attr
            
            # Extract main text content
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            text_content = soup.get_text()
            # Clean up whitespace
            lines = (line.strip() for line in text_content.splitlines())
            text_content = '\n'.join(line for line in lines if line)
            
            # Extract links
            links = []
            for link in soup.find_all('a', href=True):
                links.append({
                    'text': link.get_text().strip(),
                    'url': link['href'],
                    'title': link.get('title', '')
                })
            
            # Extract tables
            tables = []
            for table_idx, table in enumerate(soup.find_all('table')):
                rows = []
                for row in table.find_all('tr'):
                    cells = [cell.get_text().strip() for cell in row.find_all(['td', 'th'])]
                    if cells:
                        rows.append(cells)
                
                if rows:
                    tables.append({
                        'table_index': table_idx,
                        'data': rows,
                        'headers': rows[0] if rows else [],
                        'rows': rows[1:] if len(rows) > 1 else []
                    })
            
            # Extract images
            images = []
            for img in soup.find_all('img'):
                images.append({
                    'src': img.get('src', ''),
                    'alt': img.get('alt', ''),
                    'title': img.get('title', '')
                })
            
            # Structure information
            structure = {
                'headings': [{'level': tag.name, 'text': tag.get_text().strip()} 
                           for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])],
                'paragraphs_count': len(soup.find_all('p')),
                'links_count': len(links),
                'images_count': len(images)
            }
            
        except Exception as e:
            return ParsedDocument(
                content="",
                title=file_path.stem,
                metadata={'error': str(e)},
                structure={},
                images=[],
                tables=[],
                links=[],
                raw_format='html',
                encoding=encoding,
                parse_errors=[str(e)]
            )
        
        return ParsedDocument(
            content=text_content,
            title=title_text,
            metadata=metadata,
            structure=structure,
            images=images,
            tables=tables,
            links=links,
            raw_format='html',
            encoding=encoding,
            parse_errors=[]
        )

    async def _parse_csv(self, file_path: Path, content: Optional[bytes], encoding: str) -> ParsedDocument:
        """Parse CSV files as structured data."""
        
        try:
            if content:
                import io
                csv_content = content.decode(encoding)
                df = pd.read_csv(io.StringIO(csv_content))
            else:
                df = pd.read_csv(file_path, encoding=encoding)
            
            # Convert to text representation
            text_content = df.to_string(index=False)
            
            # Extract as table
            tables = [{
                'table_index': 0,
                'data': [df.columns.tolist()] + df.values.tolist(),
                'headers': df.columns.tolist(),
                'rows': df.values.tolist(),
                'shape': df.shape
            }]
            
            metadata = {
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': df.columns.tolist(),
                'data_types': df.dtypes.to_dict()
            }
            
            structure = {
                'type': 'tabular_data',
                'shape': df.shape,
                'columns': df.columns.tolist()
            }
            
        except Exception as e:
            return ParsedDocument(
                content="",
                title=file_path.stem,
                metadata={'error': str(e)},
                structure={},
                images=[],
                tables=[],
                links=[],
                raw_format='csv',
                encoding=encoding,
                parse_errors=[str(e)]
            )
        
        return ParsedDocument(
            content=text_content,
            title=file_path.stem,
            metadata=metadata,
            structure=structure,
            images=[],
            tables=tables,
            links=[],
            raw_format='csv',
            encoding=encoding,
            parse_errors=[]
        )

    async def _parse_xml(self, file_path: Path, content: Optional[bytes], encoding: str) -> ParsedDocument:
        """Parse XML documents with structure preservation."""
        
        try:
            if content:
                xml_content = content.decode(encoding)
            else:
                with open(file_path, 'r', encoding=encoding) as f:
                    xml_content = f.read()
            
            root = ET.fromstring(xml_content)
            
            # Extract text content
            text_content = ET.tostring(root, encoding='unicode', method='text')
            
            # Extract structure
            def extract_structure(element, level=0):
                structure = {
                    'tag': element.tag,
                    'text': element.text.strip() if element.text else '',
                    'attributes': element.attrib,
                    'level': level,
                    'children': []
                }
                
                for child in element:
                    structure['children'].append(extract_structure(child, level + 1))
                
                return structure
            
            structure = {
                'root': extract_structure(root),
                'total_elements': len(list(root.iter()))
            }
            
            metadata = {
                'root_tag': root.tag,
                'total_elements': len(list(root.iter())),
                'namespaces': dict([node for _, node in ET.iterparse(io.StringIO(xml_content), events=['start-ns'])])
            }
            
        except Exception as e:
            return ParsedDocument(
                content="",
                title=file_path.stem,
                metadata={'error': str(e)},
                structure={},
                images=[],
                tables=[],
                links=[],
                raw_format='xml',
                encoding=encoding,
                parse_errors=[str(e)]
            )
        
        return ParsedDocument(
            content=text_content,
            title=file_path.stem,
            metadata=metadata,
            structure=structure,
            images=[],
            tables=[],
            links=[],
            raw_format='xml',
            encoding=encoding,
            parse_errors=[]
        )

    async def _parse_text(self, file_path: Path, content: Optional[bytes], encoding: str) -> ParsedDocument:
        """Parse plain text files."""
        
        try:
            if content:
                text_content = content.decode(encoding)
            else:
                with open(file_path, 'r', encoding=encoding) as f:
                    text_content = f.read()
            
            # Basic structure analysis
            lines = text_content.split('\n')
            structure = {
                'lines_count': len(lines),
                'characters_count': len(text_content),
                'words_count': len(text_content.split()),
                'empty_lines': sum(1 for line in lines if not line.strip())
            }
            
            metadata = {
                'lines': len(lines),
                'characters': len(text_content),
                'words': len(text_content.split())
            }
            
        except Exception as e:
            return ParsedDocument(
                content="",
                title=file_path.stem,
                metadata={'error': str(e)},
                structure={},
                images=[],
                tables=[],
                links=[],
                raw_format='txt',
                encoding=encoding,
                parse_errors=[str(e)]
            )
        
        return ParsedDocument(
            content=text_content,
            title=file_path.stem,
            metadata=metadata,
            structure=structure,
            images=[],
            tables=[],
            links=[],
            raw_format='txt',
            encoding=encoding,
            parse_errors=[]
        )

    async def _parse_json(self, file_path: Path, content: Optional[bytes], encoding: str) -> ParsedDocument:
        """Parse JSON files."""
        
        try:
            if content:
                json_content = content.decode(encoding)
            else:
                with open(file_path, 'r', encoding=encoding) as f:
                    json_content = f.read()
            
            data = json.loads(json_content)
            
            # Convert to readable text
            text_content = json.dumps(data, indent=2, ensure_ascii=False)
            
            # Analyze structure
            def analyze_json_structure(obj, path="root"):
                if isinstance(obj, dict):
                    return {
                        'type': 'object',
                        'keys': list(obj.keys()),
                        'size': len(obj),
                        'children': {k: analyze_json_structure(v, f"{path}.{k}") for k, v in obj.items()}
                    }
                elif isinstance(obj, list):
                    return {
                        'type': 'array',
                        'size': len(obj),
                        'item_types': list(set(type(item).__name__ for item in obj))
                    }
                else:
                    return {
                        'type': type(obj).__name__,
                        'value': str(obj)[:100]  # Truncate long values
                    }
            
            structure = analyze_json_structure(data)
            
            metadata = {
                'json_type': type(data).__name__,
                'size': len(json_content),
                'structure_depth': json.dumps(data).count('{') + json.dumps(data).count('[')
            }
            
        except Exception as e:
            return ParsedDocument(
                content="",
                title=file_path.stem,
                metadata={'error': str(e)},
                structure={},
                images=[],
                tables=[],
                links=[],
                raw_format='json',
                encoding=encoding,
                parse_errors=[str(e)]
            )
        
        return ParsedDocument(
            content=text_content,
            title=file_path.stem,
            metadata=metadata,
            structure=structure,
            images=[],
            tables=[],
            links=[],
            raw_format='json',
            encoding=encoding,
            parse_errors=[]
        )

    async def _parse_markdown(self, file_path: Path, content: Optional[bytes], encoding: str) -> ParsedDocument:
        """Parse Markdown files with structure extraction."""
        
        try:
            if content:
                md_content = content.decode(encoding)
            else:
                with open(file_path, 'r', encoding=encoding) as f:
                    md_content = f.read()
            
            # Extract headings and structure
            import re
            
            # Find headings
            heading_pattern = r'^(#{1,6})\s+(.+)$'
            headings = []
            
            for match in re.finditer(heading_pattern, md_content, re.MULTILINE):
                level = len(match.group(1))
                text = match.group(2).strip()
                headings.append({'level': level, 'text': text})
            
            # Find links
            link_pattern = r'\[([^\]]+)\]\(([^)]+)\)'
            links = []
            
            for match in re.finditer(link_pattern, md_content):
                links.append({
                    'text': match.group(1),
                    'url': match.group(2),
                    'title': ''
                })
            
            # Find code blocks
            code_pattern = r'```(\w+)?\n(.*?)\n```'
            code_blocks = []
            
            for match in re.finditer(code_pattern, md_content, re.DOTALL):
                language = match.group(1) or 'text'
                code = match.group(2)
                code_blocks.append({
                    'language': language,
                    'code': code
                })
            
            structure = {
                'headings': headings,
                'code_blocks': code_blocks,
                'links_count': len(links),
                'lines_count': len(md_content.split('\n'))
            }
            
            metadata = {
                'headings_count': len(headings),
                'links_count': len(links),
                'code_blocks_count': len(code_blocks),
                'characters': len(md_content)
            }
            
        except Exception as e:
            return ParsedDocument(
                content="",
                title=file_path.stem,
                metadata={'error': str(e)},
                structure={},
                images=[],
                tables=[],
                links=[],
                raw_format='md',
                encoding=encoding,
                parse_errors=[str(e)]
            )
        
        return ParsedDocument(
            content=md_content,
            title=file_path.stem,
            metadata=metadata,
            structure=structure,
            images=[],
            tables=[],
            links=links,
            raw_format='md',
            encoding=encoding,
            parse_errors=[]
        )